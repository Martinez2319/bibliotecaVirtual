from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# =============================================
# MANUAL DE USUARIO
# =============================================
doc = Document()

# Título
title = doc.add_heading('Biblioteca Virtual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Manual de Usuario')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Índice
doc.add_heading('Índice', level=1)
doc.add_paragraph('1. Introducción')
doc.add_paragraph('2. Requisitos')
doc.add_paragraph('3. Acceso al Sistema')
doc.add_paragraph('4. Navegación Principal')
doc.add_paragraph('5. Catálogo de Libros')
doc.add_paragraph('6. Lectura de Libros')
doc.add_paragraph('7. Panel de Administración')
doc.add_paragraph('8. Conexión con PC Remoto')
doc.add_paragraph('9. Chatbot Asistente')
doc.add_paragraph('10. Preguntas Frecuentes')

doc.add_page_break()

# 1. Introducción
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'Biblioteca Virtual es una plataforma web que permite gestionar y leer libros digitales. '
    'El sistema permite a los usuarios explorar un catálogo de libros, leerlos en línea (PDF o texto), '
    'y a los administradores gestionar todo el contenido.'
)
doc.add_paragraph(
    'Una característica especial es la posibilidad de conectar un PC remoto para servir libros '
    'almacenados localmente, usando túneles de Cloudflare para la comunicación.'
)

# 2. Requisitos
doc.add_heading('2. Requisitos', level=1)
doc.add_heading('Para usuarios:', level=2)
doc.add_paragraph('• Navegador web moderno (Chrome, Firefox, Edge, Safari)')
doc.add_paragraph('• Conexión a internet')

doc.add_heading('Para administradores:', level=2)
doc.add_paragraph('• Credenciales de administrador')
doc.add_paragraph('• Para PC remoto: Cloudflared instalado')

# 3. Acceso al Sistema
doc.add_heading('3. Acceso al Sistema', level=1)
doc.add_heading('Registro de nuevo usuario:', level=2)
doc.add_paragraph('1. Haz clic en \"Registrarse\" en la barra de navegación')
doc.add_paragraph('2. Completa el formulario con tu nombre, email y contraseña')
doc.add_paragraph('3. Haz clic en \"Crear Cuenta\"')

doc.add_heading('Inicio de sesión:', level=2)
doc.add_paragraph('1. Haz clic en \"Iniciar Sesión\"')
doc.add_paragraph('2. Ingresa tu email y contraseña')
doc.add_paragraph('3. Haz clic en \"Entrar\"')

doc.add_heading('Credenciales de demostración:', level=2)
doc.add_paragraph('• Email: admin@biblioteca.com')
doc.add_paragraph('• Contraseña: admin123')

# 4. Navegación Principal
doc.add_heading('4. Navegación Principal', level=1)
doc.add_paragraph('La barra de navegación contiene:')
doc.add_paragraph('• Inicio: Página principal con libros destacados y recientes')
doc.add_paragraph('• Catálogo: Lista completa de libros con filtros')
doc.add_paragraph('• Admin: Panel de administración (solo administradores)')
doc.add_paragraph('• Donar: Opción para apoyar la biblioteca')

# 5. Catálogo de Libros
doc.add_heading('5. Catálogo de Libros', level=1)
doc.add_heading('Búsqueda:', level=2)
doc.add_paragraph('Usa la barra de búsqueda para encontrar libros por título o autor.')

doc.add_heading('Filtros:', level=2)
doc.add_paragraph('• Por categoría: Selecciona una categoría del menú desplegable')
doc.add_paragraph('• Ordenar: Por fecha, popularidad o título')

doc.add_heading('Ver detalles:', level=2)
doc.add_paragraph('Haz clic en cualquier libro para ver su información completa.')

# 6. Lectura de Libros
doc.add_heading('6. Lectura de Libros', level=1)
doc.add_paragraph('Al abrir un libro, verás:')
doc.add_paragraph('• Portada del libro')
doc.add_paragraph('• Título y autor')
doc.add_paragraph('• Descripción')
doc.add_paragraph('• Formatos disponibles (PDF y/o Texto)')

doc.add_heading('Lector:', level=2)
doc.add_paragraph('• Si hay PDF y texto, puedes alternar entre ambos formatos')
doc.add_paragraph('• En modo texto, puedes ajustar el tamaño de fuente (A- / A+)')
doc.add_paragraph('• Usa \"← Volver\" para regresar al catálogo')

doc.add_heading('Límite para invitados:', level=2)
doc.add_paragraph('Los usuarios no registrados solo pueden leer 1 libro. Regístrate para acceso ilimitado.')

# 7. Panel de Administración
doc.add_heading('7. Panel de Administración', level=1)
doc.add_paragraph('Accede desde el enlace \"Admin\" en la navegación (solo administradores).')

doc.add_heading('Estadísticas:', level=2)
doc.add_paragraph('Muestra el total de libros, usuarios, categorías y estado del PC remoto.')

doc.add_heading('Gestión de Libros:', level=2)
doc.add_paragraph('• Agregar: Clic en \"+ Nuevo Libro\"')
doc.add_paragraph('• Editar: Clic en el ícono de lápiz ✏️')
doc.add_paragraph('• Eliminar: Clic en el ícono de basura 🗑️')

doc.add_heading('Campos del libro:', level=2)
doc.add_paragraph('• Título y Autor (obligatorios)')
doc.add_paragraph('• Descripción')
doc.add_paragraph('• Categorías (separadas por coma)')
doc.add_paragraph('• URL de portada (imagen)')
doc.add_paragraph('• URL del PDF')
doc.add_paragraph('• Contenido en texto')

doc.add_heading('Gestión de Categorías:', level=2)
doc.add_paragraph('Crea, edita y elimina categorías para organizar los libros.')

doc.add_heading('Gestión de Usuarios:', level=2)
doc.add_paragraph('Cambia roles (usuario/admin) o elimina usuarios.')

# 8. Conexión con PC Remoto
doc.add_heading('8. Conexión con PC Remoto', level=1)
doc.add_paragraph(
    'Esta función permite servir libros desde tu computadora personal, '
    'sin necesidad de subirlos a un servidor.'
)

doc.add_heading('Configuración:', level=2)
doc.add_paragraph('1. En tu PC, crea las carpetas:')
doc.add_paragraph('   C:\\MisLibros\\covers\\ (para portadas)')
doc.add_paragraph('   C:\\MisLibros\\pdfs\\ (para PDFs)')
doc.add_paragraph('')
doc.add_paragraph('2. Abre una terminal y ejecuta:')
doc.add_paragraph('   cloudflared tunnel --url http://localhost:3001')
doc.add_paragraph('')
doc.add_paragraph('3. En otra terminal del servidor de biblioteca:')
doc.add_paragraph('   cloudflared tunnel --url http://localhost:3000')
doc.add_paragraph('')
doc.add_paragraph('4. Ejecuta iniciar.bat en la carpeta pc-server')
doc.add_paragraph('5. Pega las URLs cuando te las pida')

doc.add_heading('Uso de archivos remotos:', level=2)
doc.add_paragraph('Al agregar un libro, usa el prefijo \"remote:\" seguido del nombre:')
doc.add_paragraph('• Portada: remote:mi-libro.jpg')
doc.add_paragraph('• PDF: remote:mi-libro.pdf')

# 9. Chatbot Asistente
doc.add_heading('9. Chatbot Asistente', level=1)
doc.add_paragraph(
    'En la esquina inferior derecha aparece el chatbot de la biblioteca. '
    'Puedes preguntarle sobre libros, categorías o cómo usar el sistema.'
)

# 10. Preguntas Frecuentes
doc.add_heading('10. Preguntas Frecuentes', level=1)

doc.add_heading('¿Cómo cambio mi contraseña?', level=2)
doc.add_paragraph('Actualmente debes contactar al administrador.')

doc.add_heading('¿Por qué no puedo leer más libros?', level=2)
doc.add_paragraph('Los invitados tienen límite de 1 libro. Regístrate para acceso ilimitado.')

doc.add_heading('¿Qué formatos de libros soporta?', level=2)
doc.add_paragraph('PDF y texto plano.')

doc.add_heading('¿Puedo usar Google Drive para las portadas?', level=2)
doc.add_paragraph('Sí, pega el enlace completo del archivo compartido.')

doc.save('docs/Manual_Usuario.docx')
print('✅ Manual de Usuario creado')
