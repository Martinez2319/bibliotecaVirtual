from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================
# DOCUMENTACIÓN DEL CÓDIGO
# =============================================
doc = Document()

# Título
title = doc.add_heading('Biblioteca Virtual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Documentación del Código')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Índice
doc.add_heading('Índice', level=1)
doc.add_paragraph('1. server.js - Servidor Principal')
doc.add_paragraph('2. Modelos de MongoDB')
doc.add_paragraph('3. Rutas de la API')
doc.add_paragraph('4. Frontend - HTML')
doc.add_paragraph('5. Frontend - CSS')
doc.add_paragraph('6. Frontend - JavaScript')
doc.add_paragraph('7. PC Server')

doc.add_page_break()

# ============================================
# 1. SERVER.JS
# ============================================
doc.add_heading('1. server.js - Servidor Principal', level=1)
doc.add_paragraph('Este es el archivo principal que inicia el servidor Express.')

doc.add_heading('1.1 Importaciones', level=2)
code = '''require('dotenv').config();           // Carga variables de .env
const express = require('express');      // Framework web
const mongoose = require('mongoose');    // ODM para MongoDB
const cors = require('cors');            // Habilita CORS
const cookieParser = require('cookie-parser'); // Parse de cookies
const path = require('path');            // Manejo de rutas
const jwt = require('jsonwebtoken');     // Tokens JWT
const bcrypt = require('bcryptjs');      // Hash de contraseñas'''
doc.add_paragraph(code)

doc.add_heading('1.2 Importación de Modelos', level=2)
code = '''const User = require('./models/User');
const Book = require('./models/Book');
const Category = require('./models/Category');'''
doc.add_paragraph(code)
doc.add_paragraph('Los modelos definen la estructura de los documentos en MongoDB.')

doc.add_heading('1.3 Importación de Rutas', level=2)
code = '''const authRoutes = require('./routes/auth');
const booksRoutes = require('./routes/books');
const categoriesRoutes = require('./routes/categories');
const usersRoutes = require('./routes/users');
const paypalRoutes = require('./routes/paypal');
const remoteRoutes = require('./routes/remote');'''
doc.add_paragraph(code)
doc.add_paragraph('Cada archivo de rutas contiene los endpoints de una funcionalidad.')

doc.add_heading('1.4 Configuración de Express', level=2)
code = '''const app = express();

app.use(cors({ origin: true, credentials: true }));
app.use(express.json({ limit: '10mb' }));
app.use(cookieParser());
app.use(express.static(path.join(__dirname, 'public')));'''
doc.add_paragraph(code)
doc.add_paragraph('• cors: Permite peticiones desde cualquier origen con cookies')
doc.add_paragraph('• express.json: Parsea body JSON con límite de 10MB')
doc.add_paragraph('• cookieParser: Permite leer cookies de las peticiones')
doc.add_paragraph('• express.static: Sirve archivos estáticos desde /public')

doc.add_heading('1.5 Middleware de Autenticación', level=2)
code = '''app.use(async (req, res, next) => {
  const token = req.cookies.token || 
                req.headers.authorization?.split(' ')[1];
  if (token) {
    try {
      const decoded = jwt.verify(token, process.env.JWT_SECRET);
      req.user = await User.findById(decoded.id)
                           .select('-passwordHash');
    } catch (e) {
      req.user = null;
    }
  }
  next();
});'''
doc.add_paragraph(code)
doc.add_paragraph('Este middleware se ejecuta en CADA petición:')
doc.add_paragraph('1. Busca token en cookies o header Authorization')
doc.add_paragraph('2. Si existe, verifica la firma con JWT_SECRET')
doc.add_paragraph('3. Decodifica el ID del usuario')
doc.add_paragraph('4. Busca el usuario en MongoDB (sin la contraseña)')
doc.add_paragraph('5. Asigna el usuario a req.user')
doc.add_paragraph('6. Si falla, req.user queda null')

doc.add_heading('1.6 Registro de Rutas', level=2)
code = '''app.use('/api/auth', authRoutes);
app.use('/api/books', booksRoutes);
app.use('/api/categories', categoriesRoutes);
app.use('/api/users', usersRoutes);
app.use('/api/paypal', paypalRoutes);
app.use('/api/remote', remoteRoutes);'''
doc.add_paragraph(code)
doc.add_paragraph('Cada app.use() monta las rutas en un prefijo específico.')

doc.add_heading('1.7 Ruta de Estadísticas', level=2)
code = '''app.get('/api/stats', async (req, res) => {
  try {
    if (!req.user || req.user.role !== 'admin') {
      return res.status(403).json({ error: 'Acceso denegado' });
    }
    const [totalBooks, totalUsers, totalCategories] = 
      await Promise.all([
        Book.countDocuments(),
        User.countDocuments(),
        Category.countDocuments()
      ]);
    res.json({ totalBooks, totalUsers, totalCategories });
  } catch (error) {
    res.status(500).json({ error: 'Error al obtener estadísticas' });
  }
});'''
doc.add_paragraph(code)
doc.add_paragraph('• Verifica que el usuario sea admin')
doc.add_paragraph('• Usa Promise.all para ejecutar 3 consultas en paralelo')
doc.add_paragraph('• Retorna los conteos en un objeto JSON')

doc.add_heading('1.8 Rutas HTML', level=2)
code = '''app.get('/', (req, res) => 
  res.sendFile(path.join(__dirname, 'public/html/index.html')));
app.get('/catalog', (req, res) => 
  res.sendFile(path.join(__dirname, 'public/html/catalog.html')));
// ... más rutas'''
doc.add_paragraph(code)
doc.add_paragraph('Cada ruta sirve un archivo HTML estático.')

doc.add_heading('1.9 Manejo de Errores', level=2)
code = '''// 404 - Ruta no encontrada
app.use((req, res) => {
  res.status(404).json({ error: 'Ruta no encontrada' });
});

// Error global
app.use((err, req, res, next) => {
  console.error('Error:', err.message);
  res.status(500).json({ error: 'Error interno del servidor' });
});'''
doc.add_paragraph(code)
doc.add_paragraph('• El middleware 404 captura rutas no definidas')
doc.add_paragraph('• El middleware de error captura excepciones no manejadas')

doc.add_heading('1.10 Conexión a MongoDB e Inicio', level=2)
code = '''const PORT = process.env.PORT || 3000;

mongoose.connect(`${process.env.MONGO_URL}/${process.env.DB_NAME}`)
  .then(() => {
    console.log('✅ MongoDB conectado');
    app.listen(PORT, () => 
      console.log(`🚀 Servidor en http://localhost:${PORT}`));
  })
  .catch(err => {
    console.error('❌ Error de conexión:', err.message);
    process.exit(1);
  });'''
doc.add_paragraph(code)
doc.add_paragraph('• Construye la URL de MongoDB concatenando MONGO_URL y DB_NAME')
doc.add_paragraph('• Si conecta, inicia el servidor Express')
doc.add_paragraph('• Si falla, termina el proceso con código 1')

doc.add_page_break()

# ============================================
# 2. MODELOS
# ============================================
doc.add_heading('2. Modelos de MongoDB', level=1)
doc.add_paragraph('Los modelos definen la estructura de los documentos usando Mongoose.')

doc.add_heading('2.1 User.js', level=2)
code = '''const mongoose = require('mongoose');

const userSchema = new mongoose.Schema({
  name: { type: String, required: true },
  email: { type: String, required: true, unique: true },
  passwordHash: { type: String, required: true },
  role: { type: String, enum: ['user', 'admin'], default: 'user' },
  createdAt: { type: Date, default: Date.now }
});

module.exports = mongoose.model('User', userSchema);'''
doc.add_paragraph(code)
doc.add_paragraph('• name: Nombre del usuario')
doc.add_paragraph('• email: Único, usado para login')
doc.add_paragraph('• passwordHash: Contraseña hasheada con bcrypt')
doc.add_paragraph('• role: user o admin (default: user)')
doc.add_paragraph('• createdAt: Fecha de registro')

doc.add_heading('2.2 Book.js', level=2)
code = '''const bookSchema = new mongoose.Schema({
  title: { type: String, required: true },
  author: { type: String, required: true },
  description: String,
  categories: [String],
  coverUrl: String,
  pdfUrl: String,
  content: String,
  views: { type: Number, default: 0 },
  createdAt: { type: Date, default: Date.now }
});'''
doc.add_paragraph(code)
doc.add_paragraph('• title/author: Datos básicos del libro')
doc.add_paragraph('• categories: Array de strings con nombres de categorías')
doc.add_paragraph('• coverUrl: URL de la imagen de portada')
doc.add_paragraph('• pdfUrl: URL del archivo PDF')
doc.add_paragraph('• content: Texto del libro (alternativa al PDF)')
doc.add_paragraph('• views: Contador de vistas')

doc.add_heading('2.3 Category.js', level=2)
code = '''const categorySchema = new mongoose.Schema({
  name: { type: String, required: true },
  slug: { type: String, required: true, unique: true },
  description: String
});'''
doc.add_paragraph(code)
doc.add_paragraph('• name: Nombre visible de la categoría')
doc.add_paragraph('• slug: URL-friendly (ej: \"ciencia-ficcion\")')
doc.add_paragraph('• description: Descripción opcional')

doc.add_heading('2.4 RemoteSource.js', level=2)
code = '''const remoteSourceSchema = new mongoose.Schema({
  name: { type: String, default: 'Mi PC' },
  url: { type: String, required: true },
  apiKey: { type: String, required: true },
  isOnline: { type: Boolean, default: false },
  lastSeen: { type: Date, default: null },
  createdAt: { type: Date, default: Date.now }
});'''
doc.add_paragraph(code)
doc.add_paragraph('• url: URL pública del túnel Cloudflare')
doc.add_paragraph('• apiKey: Clave de autenticación')
doc.add_paragraph('• isOnline: Estado de conexión')
doc.add_paragraph('• lastSeen: Último heartbeat recibido')

doc.add_heading('2.5 AccessLog.js', level=2)
code = '''const accessLogSchema = new mongoose.Schema({
  identifier: { type: String, required: true },
  bookId: { type: String, required: true },
  createdAt: { type: Date, default: Date.now }
});'''
doc.add_paragraph(code)
doc.add_paragraph('Controla el límite de lecturas para invitados.')
doc.add_paragraph('• identifier: IP del usuario')
doc.add_paragraph('• bookId: ID del libro leído')

doc.add_page_break()

# ============================================
# 3. RUTAS
# ============================================
doc.add_heading('3. Rutas de la API', level=1)

doc.add_heading('3.1 routes/auth.js', level=2)
doc.add_paragraph('Maneja registro, login y logout de usuarios.')

code = '''// POST /api/auth/register
router.post('/register', async (req, res) => {
  const { name, email, password } = req.body;
  
  // Verificar si email ya existe
  const exists = await User.findOne({ email });
  if (exists) return res.status(400).json({ error: 'Email ya registrado' });
  
  // Hashear contraseña
  const passwordHash = await bcrypt.hash(password, 10);
  
  // Crear usuario
  const user = await User.create({ name, email, passwordHash });
  
  // Generar token
  const token = jwt.sign({ id: user._id }, process.env.JWT_SECRET);
  
  // Guardar en cookie
  res.cookie('token', token, { httpOnly: true });
  res.json({ user: { name, email, role: user.role } });
});'''
doc.add_paragraph(code)

code = '''// POST /api/auth/login
router.post('/login', async (req, res) => {
  const { email, password } = req.body;
  
  const user = await User.findOne({ email });
  if (!user) return res.status(401).json({ error: 'Credenciales inválidas' });
  
  const valid = await bcrypt.compare(password, user.passwordHash);
  if (!valid) return res.status(401).json({ error: 'Credenciales inválidas' });
  
  const token = jwt.sign({ id: user._id }, process.env.JWT_SECRET);
  res.cookie('token', token, { httpOnly: true });
  res.json({ user: { name: user.name, email, role: user.role } });
});'''
doc.add_paragraph(code)

doc.add_heading('3.2 routes/books.js', level=2)
doc.add_paragraph('CRUD completo de libros con soporte para URLs remotas.')

code = '''// Helper: Convertir URLs remote: a URLs del proxy
async function processBookUrls(book) {
  const bookObj = book.toObject ? book.toObject() : { ...book };
  
  if (bookObj.coverUrl?.startsWith('remote:')) {
    const filename = bookObj.coverUrl.replace('remote:', '');
    bookObj.coverUrl = `/api/remote/file/cover/${encodeURIComponent(filename)}`;
  }
  
  if (bookObj.pdfUrl?.startsWith('remote:')) {
    const filename = bookObj.pdfUrl.replace('remote:', '');
    bookObj.pdfUrl = `/api/remote/file/pdf/${encodeURIComponent(filename)}`;
  }
  
  return bookObj;
}'''
doc.add_paragraph(code)
doc.add_paragraph('Esta función transforma URLs con prefijo \"remote:\" en URLs del proxy.')

code = '''// GET /api/books
router.get('/', async (req, res) => {
  const { search, category, sort = 'createdAt' } = req.query;
  let query = {};
  
  if (search) {
    query.$or = [
      { title: { $regex: search, $options: 'i' } },
      { author: { $regex: search, $options: 'i' } }
    ];
  }
  if (category) query.categories = category;
  
  const books = await Book.find(query)
    .sort({ [sort]: sort === 'title' ? 1 : -1 });
  const processedBooks = await processBooksArray(books);
  res.json(processedBooks);
});'''
doc.add_paragraph(code)
doc.add_paragraph('• Soporta búsqueda por título o autor (case-insensitive)')
doc.add_paragraph('• Filtra por categoría')
doc.add_paragraph('• Ordena por el campo especificado')

doc.add_heading('3.3 routes/remote.js', level=2)
doc.add_paragraph('Maneja la comunicación con el PC remoto.')

code = '''// Middleware: verificar API key
const verifyServerKey = (req, res, next) => {
  const key = req.headers['x-api-key'] || req.query.apiKey;
  if (key !== process.env.REMOTE_API_KEY) {
    return res.status(401).json({ error: 'API key inválida' });
  }
  next();
};'''
doc.add_paragraph(code)

code = '''// POST /api/remote/register - PC se registra
router.post('/register', verifyServerKey, async (req, res) => {
  const { url, name } = req.body;
  
  let source = await RemoteSource.findOne({ 
    apiKey: process.env.REMOTE_API_KEY 
  });
  
  if (source) {
    source.url = url;
    source.isOnline = true;
    source.lastSeen = new Date();
    await source.save();
  } else {
    source = await RemoteSource.create({
      url, name, 
      apiKey: process.env.REMOTE_API_KEY,
      isOnline: true
    });
  }
  
  res.json({ success: true });
});'''
doc.add_paragraph(code)

code = '''// GET /api/remote/file/:type/:filename - Proxy de archivos
router.get('/file/:type/:filename', async (req, res) => {
  const { type, filename } = req.params;
  
  const source = await RemoteSource.findOne({ 
    apiKey: process.env.REMOTE_API_KEY, 
    isOnline: true 
  });
  
  if (!source) {
    return res.status(404).json({ error: 'PC no conectado' });
  }
  
  const response = await axios.get(
    `${source.url}/file/${type}/${filename}`,
    {
      headers: { 'x-api-key': process.env.REMOTE_API_KEY },
      responseType: 'stream'
    }
  );
  
  response.data.pipe(res);
});'''
doc.add_paragraph(code)
doc.add_paragraph('Este endpoint actúa como proxy, redirigiendo la petición al PC remoto.')

doc.add_page_break()

# ============================================
# 4. FRONTEND HTML
# ============================================
doc.add_heading('4. Frontend - HTML', level=1)

doc.add_heading('4.1 Estructura común', level=2)
doc.add_paragraph('Todas las páginas siguen esta estructura:')
code = '''<!DOCTYPE html>
<html lang=\"es\">
<head>
  <meta charset=\"UTF-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
  <title>Título - Biblioteca Virtual</title>
  <link rel=\"stylesheet\" href=\"/css/index.css\">
  <!-- CSS específico de la página -->
</head>
<body>
  <nav class=\"navbar\">...</nav>
  
  <!-- Contenido de la página -->
  
  <script src=\"/javascript/archivo.js\"></script>
  
  <!-- Chatbot Botpress -->
  <script src=\"https://cdn.botpress.cloud/webchat/v3.6/inject.js\"></script>
</body>
</html>'''
doc.add_paragraph(code)

doc.add_heading('4.2 index.html - Página principal', level=2)
doc.add_paragraph('Contiene:')
doc.add_paragraph('• Navbar con navegación')
doc.add_paragraph('• Hero con imagen de fondo')
doc.add_paragraph('• Carruseles de libros destacados y recientes')
doc.add_paragraph('• Grid de categorías')
doc.add_paragraph('• Modal de donación')

doc.add_heading('4.3 catalog.html - Catálogo', level=2)
doc.add_paragraph('Contiene:')
doc.add_paragraph('• Barra de búsqueda')
doc.add_paragraph('• Filtros de categoría y ordenamiento')
doc.add_paragraph('• Grid de libros dinámico')

doc.add_heading('4.4 dashboard.html - Panel admin', level=2)
doc.add_paragraph('Contiene:')
doc.add_paragraph('• Cards de estadísticas')
doc.add_paragraph('• Tabs para libros, categorías, usuarios y PC remoto')
doc.add_paragraph('• Tablas con acciones CRUD')
doc.add_paragraph('• Modales para crear/editar')

doc.add_heading('4.5 reader.html - Lector de libros', level=2)
doc.add_paragraph('Contiene:')
doc.add_paragraph('• Header con título y controles')
doc.add_paragraph('• Toggle PDF/Texto (si ambos disponibles)')
doc.add_paragraph('• Iframe para PDF')
doc.add_paragraph('• Div para contenido de texto')

doc.add_page_break()

# ============================================
# 5. FRONTEND CSS
# ============================================
doc.add_heading('5. Frontend - CSS', level=1)

doc.add_heading('5.1 Variables CSS (index.css)', level=2)
code = ''':root {
  --primary: #1a1a2e;   /* Azul oscuro */
  --accent: #e94560;    /* Rosa/rojo */
  --bg: #fafafa;        /* Fondo gris claro */
  --card: #ffffff;      /* Fondo de cards */
  --text: #1a1a2e;      /* Texto oscuro */
  --muted: #6b7280;     /* Texto secundario */
  --border: #e5e7eb;    /* Bordes */
}'''
doc.add_paragraph(code)

doc.add_heading('5.2 Sistema de Grid Responsivo', level=2)
code = '''.books-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
  gap: 1rem;
}

@media (max-width: 480px) {
  .books-grid {
    grid-template-columns: repeat(2, 1fr);
  }
}'''
doc.add_paragraph(code)
doc.add_paragraph('• auto-fill: Crea tantas columnas como quepan')
doc.add_paragraph('• minmax: Mínimo 150px, máximo 1fr')
doc.add_paragraph('• En móvil: Fuerza 2 columnas')

doc.add_heading('5.3 Tipografía Fluida con clamp()', level=2)
code = '''.hero h1 {
  font-size: clamp(1.5rem, 5vw, 2.5rem);
}'''
doc.add_paragraph(code)
doc.add_paragraph('• Mínimo: 1.5rem')
doc.add_paragraph('• Preferido: 5% del viewport width')
doc.add_paragraph('• Máximo: 2.5rem')

doc.add_heading('5.4 Breakpoints', level=2)
doc.add_paragraph('El sistema usa dos breakpoints principales:')
doc.add_paragraph('• 768px: Tablets')
doc.add_paragraph('• 480px: Móviles')

doc.add_page_break()

# ============================================
# 6. FRONTEND JS
# ============================================
doc.add_heading('6. Frontend - JavaScript', level=1)

doc.add_heading('6.1 index.js - Página principal', level=2)
code = '''// Verificar autenticación
async function checkAuth() {
  const res = await fetch('/api/auth/me');
  if (res.ok) {
    const user = await res.json();
    // Mostrar menú de usuario, ocultar login
    document.getElementById('authLinks').style.display = 'none';
    document.getElementById('userMenu').style.display = 'inline';
    document.getElementById('userName').textContent = user.name;
    if (user.role === 'admin') {
      document.getElementById('adminLink').style.display = 'inline';
    }
  }
}'''
doc.add_paragraph(code)

code = '''// Cargar libros destacados
async function loadFeatured() {
  const res = await fetch('/api/books/featured');
  const books = await res.json();
  
  document.getElementById('featuredBooks').innerHTML = books.map(b => `
    <div class=\"book-card\" onclick=\"location.href='/book/${b._id}'\">
      <img src=\"${b.coverUrl}\" alt=\"${b.title}\">
      <div class=\"book-card-content\">
        <h3>${b.title}</h3>
        <p>${b.author}</p>
      </div>
    </div>
  `).join('');
}'''
doc.add_paragraph(code)

doc.add_heading('6.2 dashboard.js - Panel admin', level=2)
code = '''// Cargar estado del PC remoto
async function loadRemoteStatus() {
  const res = await fetch('/api/remote/status');
  const data = await res.json();
  
  if (data.connected) {
    // Mostrar como online
    pcStatusValue.textContent = '✅ Online';
    pcStatusCard.style.background = '#dcfce7';
    // Cargar lista de archivos
    loadRemoteFiles();
  } else {
    // Mostrar como offline
    pcStatusValue.textContent = '❌ Offline';
  }
}'''
doc.add_paragraph(code)

code = '''// Guardar libro
document.getElementById('bookForm').addEventListener('submit', async (e) => {
  e.preventDefault();
  
  const id = document.getElementById('bookId').value;
  const data = {
    title: document.getElementById('bookTitle').value,
    author: document.getElementById('bookAuthor').value,
    description: document.getElementById('bookDesc').value,
    categories: document.getElementById('bookCategories').value
      .split(',').map(c => c.trim()).filter(c => c),
    coverUrl: document.getElementById('bookCover').value,
    pdfUrl: document.getElementById('bookPdf').value,
    content: document.getElementById('bookContent').value
  };
  
  await fetch(id ? `/api/books/${id}` : '/api/books', {
    method: id ? 'PUT' : 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify(data)
  });
  
  closeBookModal();
  loadBooks();
});'''
doc.add_paragraph(code)

doc.add_page_break()

# ============================================
# 7. PC SERVER
# ============================================
doc.add_heading('7. PC Server', level=1)

doc.add_heading('7.1 server.js del PC', level=2)
code = '''// Configuración desde config.json
const config = require('./config.json');

// Carpetas de archivos
const BOOKS_FOLDER = config.booksFolder || 'C:\\\\MisLibros';
const PDF_FOLDER = path.join(BOOKS_FOLDER, 'pdfs');
const COVERS_FOLDER = path.join(BOOKS_FOLDER, 'covers');'''
doc.add_paragraph(code)

code = '''// Verificar API Key en cada petición
const verifyApiKey = (req, res, next) => {
  const key = req.headers['x-api-key'];
  if (key !== config.apiKey) {
    return res.status(401).json({ error: 'API key inválida' });
  }
  next();
};'''
doc.add_paragraph(code)

code = '''// Listar archivos disponibles
app.get('/files', verifyApiKey, (req, res) => {
  const pdfs = fs.readdirSync(PDF_FOLDER)
    .filter(f => f.endsWith('.pdf'))
    .map(f => ({
      name: f,
      size: fs.statSync(path.join(PDF_FOLDER, f)).size
    }));
  
  const covers = fs.readdirSync(COVERS_FOLDER)
    .filter(f => /\\.(jpg|png)$/i.test(f))
    .map(f => ({
      name: f,
      size: fs.statSync(path.join(COVERS_FOLDER, f)).size
    }));
  
  res.json({ pdfs, covers, booksFolder: BOOKS_FOLDER });
});'''
doc.add_paragraph(code)

code = '''// Servir archivo PDF
app.get('/file/pdf/:filename', verifyApiKey, (req, res) => {
  const filepath = path.join(PDF_FOLDER, req.params.filename);
  
  // Validar que el path esté dentro de la carpeta
  if (!filepath.startsWith(PDF_FOLDER)) {
    return res.status(400).json({ error: 'Path inválido' });
  }
  
  if (!fs.existsSync(filepath)) {
    return res.status(404).json({ error: 'No encontrado' });
  }
  
  res.setHeader('Content-Type', 'application/pdf');
  fs.createReadStream(filepath).pipe(res);
});'''
doc.add_paragraph(code)

code = '''// Registrarse con el servidor principal
async function registerWithServer(url) {
  await axios.post(`${config.serverUrl}/api/remote/register`, {
    url: url,
    name: config.pcName
  }, {
    headers: { 'x-api-key': config.apiKey }
  });
}

// Heartbeat cada 30 segundos
setInterval(async () => {
  await axios.post(`${config.serverUrl}/api/remote/heartbeat`, {}, {
    headers: { 'x-api-key': config.apiKey }
  });
}, 30000);'''
doc.add_paragraph(code)

doc.add_heading('7.2 iniciar.bat', level=2)
code = '''@echo off
set /p TUNNEL_URL=\"URL del PC-SERVER: \"
set /p BIBLIOTECA_URL=\"URL de BIBLIOTECA-VIRTUAL: \"

:: Actualizar config.json con la nueva URL
echo { > config.json
echo   \"serverUrl\": \"%BIBLIOTECA_URL%\",
echo   \"apiKey\": \"...\",
echo   ...
echo } >> config.json

set \"TUNNEL_URL=%TUNNEL_URL%\"
node server.js'''
doc.add_paragraph(code)
doc.add_paragraph('El .bat pide las URLs de Cloudflare, actualiza config.json y inicia el servidor.')

doc.save('docs/Documentacion_Codigo.docx')
print('✅ Documentación del Código creada')
