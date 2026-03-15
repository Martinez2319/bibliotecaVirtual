from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Establece el color de fondo de una celda"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_colored_paragraph(doc, text, color_hex, bold=False, size=12):
    """Agrega un párrafo con color"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.bold = bold
    run.font.size = Pt(size)
    return p

# =============================================
# DOCUMENTO COMPLETO DEL PROYECTO
# =============================================
doc = Document()

# ============= PORTADA =============
title = doc.add_heading('📚 BIBLIOTECA VIRTUAL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(26, 26, 46)  # #1a1a2e

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Documentación Completa del Proyecto')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

# Info box
info_table = doc.add_table(rows=1, cols=1)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = info_table.rows[0].cells[0]
set_cell_shading(cell, 'E8F4FD')
cell.text = """🌐 Sistema de biblioteca digital con soporte para PC remoto
📖 Lectura de libros en PDF y texto
🤖 Asistente virtual integrado"""
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Versión y fecha
version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Versión 2.0 | Marzo 2024')
run.font.color.rgb = RGBColor(156, 163, 175)

doc.add_page_break()

# ============= ÍNDICE =============
doc.add_heading('📋 Índice', level=1)

indices = [
    ('1. Visión General del Proyecto', '4361ee'),
    ('2. Arquitectura del Sistema', '10b981'),
    ('3. Funcionalidades', '8b5cf6'),
    ('4. Limitaciones', 'ef4444'),
    ('5. Mapa de Dependencias de Archivos', 'f59e0b'),
    ('6. Flujos de Datos', '06b6d4'),
    ('7. Diagrama de Base de Datos', 'ec4899'),
    ('8. Guía de Despliegue', '6366f1'),
]

for idx, (text, color) in enumerate(indices):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(12)

doc.add_page_break()

# ============= 1. VISIÓN GENERAL =============
h = doc.add_heading('1. 🎯 Visión General del Proyecto', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(67, 97, 238)

doc.add_paragraph(
    'Biblioteca Virtual es una plataforma web completa para gestionar y compartir libros digitales. '
    'Diseñada para funcionar tanto con almacenamiento en la nube como con archivos locales desde un PC remoto.'
)

# Características principales - Tabla con colores
doc.add_heading('Características Principales', level=2)

features_table = doc.add_table(rows=5, cols=2)
features_table.style = 'Table Grid'

features = [
    ('📖 Biblioteca Digital', 'Catálogo de libros con búsqueda, filtros y categorías'),
    ('📄 Lector Integrado', 'Soporte para PDF y texto, con controles de lectura'),
    ('🖥️ PC Remoto', 'Servir archivos desde tu computadora personal'),
    ('🤖 Chatbot', 'Asistente virtual con Botpress'),
    ('👥 Multi-usuario', 'Sistema de roles (admin/usuario) con autenticación JWT'),
]

for i, (feature, desc) in enumerate(features):
    row = features_table.rows[i]
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    cell1.text = feature
    cell2.text = desc
    # Color de fondo alternado
    if i % 2 == 0:
        set_cell_shading(cell1, 'EEF2FF')
        set_cell_shading(cell2, 'EEF2FF')
    # Negrita en primera columna
    cell1.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# Stack tecnológico
doc.add_heading('Stack Tecnológico', level=2)

stack_table = doc.add_table(rows=4, cols=3)
stack_table.style = 'Table Grid'

stack = [
    ('Capa', 'Tecnología', 'Versión'),
    ('🖥️ Backend', 'Node.js + Express', 'v18+ / v4.x'),
    ('🗄️ Base de Datos', 'MongoDB', 'v6+'),
    ('🎨 Frontend', 'HTML5 + CSS3 + JS', 'ES6+'),
]

colors = ['1a1a2e', '4361ee', '10b981', 'f59e0b']

for i, (col1, col2, col3) in enumerate(stack):
    row = stack_table.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    row.cells[2].text = col3
    if i == 0:  # Header
        for cell in row.cells:
            set_cell_shading(cell, '1a1a2e')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ============= 2. ARQUITECTURA =============
h = doc.add_heading('2. 🏗️ Arquitectura del Sistema', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(16, 185, 129)

doc.add_heading('Diagrama de Arquitectura', level=2)

# Diagrama ASCII en tabla para mejor formato
arch_table = doc.add_table(rows=1, cols=1)
arch_table.style = 'Table Grid'
cell = arch_table.rows[0].cells[0]
set_cell_shading(cell, 'F0FDF4')

diagram = '''
┌─────────────────────────────────────────────────────────────────┐
│                        USUARIOS                                  │
│                    (Navegador Web)                               │
└─────────────────────────┬───────────────────────────────────────┘
                          │
                          ▼
┌─────────────────────────────────────────────────────────────────┐
│                      FRONTEND                                    │
│  ┌─────────┐  ┌─────────┐  ┌─────────┐  ┌─────────────┐        │
│  │ index   │  │ catalog │  │ reader  │  │  dashboard  │        │
│  │  .html  │  │  .html  │  │  .html  │  │    .html    │        │
│  └─────────┘  └─────────┘  └─────────┘  └─────────────┘        │
│        CSS: index.css, catalog.css, dashboard.css, etc.         │
│        JS: index.js, dashboard.js, login.js, etc.               │
└─────────────────────────┬───────────────────────────────────────┘
                          │ fetch('/api/...')
                          ▼
┌─────────────────────────────────────────────────────────────────┐
│                       BACKEND                                    │
│                     (server.js)                                  │
│  ┌──────────────────────────────────────────────────────────┐   │
│  │                    EXPRESS.JS                             │   │
│  │  ┌─────────┐ ┌─────────┐ ┌─────────┐ ┌─────────────────┐ │   │
│  │  │  auth   │ │  books  │ │ users   │ │     remote      │ │   │
│  │  │ routes  │ │ routes  │ │ routes  │ │     routes      │ │   │
│  │  └─────────┘ └─────────┘ └─────────┘ └─────────────────┘ │   │
│  └──────────────────────────────────────────────────────────┘   │
└──────────────┬──────────────────────────────┬───────────────────┘
               │                              │
               ▼                              ▼
┌──────────────────────────┐    ┌─────────────────────────────────┐
│        MONGODB           │    │         PC REMOTO               │
│   ┌─────────────────┐    │    │    (Cloudflare Tunnel)          │
│   │ users           │    │    │  ┌─────────────────────────┐    │
│   │ books           │    │    │  │  C:\MisLibros\          │    │
│   │ categories      │    │    │  │  ├── pdfs\              │    │
│   │ remotesources   │    │    │  │  └── covers\            │    │
│   └─────────────────┘    │    │  └─────────────────────────┘    │
└──────────────────────────┘    └─────────────────────────────────┘
'''
cell.text = diagram
for para in cell.paragraphs:
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

doc.add_paragraph()

# Explicación de componentes
doc.add_heading('Componentes del Sistema', level=2)

comp_table = doc.add_table(rows=5, cols=3)
comp_table.style = 'Table Grid'

components = [
    ('Componente', 'Descripción', 'Puerto'),
    ('Frontend', 'Páginas HTML estáticas servidas por Express', '3000'),
    ('Backend API', 'Servidor Express con endpoints REST', '3000'),
    ('MongoDB', 'Base de datos NoSQL (Atlas o local)', '27017'),
    ('PC Server', 'Servidor local para archivos', '3001'),
]

for i, (c1, c2, c3) in enumerate(components):
    row = comp_table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, '10b981')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ============= 3. FUNCIONALIDADES =============
h = doc.add_heading('3. ✅ Funcionalidades (Lo que PUEDE hacer)', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(139, 92, 246)

# Tabla de funcionalidades con estado
func_table = doc.add_table(rows=16, cols=3)
func_table.style = 'Table Grid'

functionalities = [
    ('Categoría', 'Funcionalidad', 'Estado'),
    ('👥 Usuarios', 'Registro de nuevos usuarios', '✅'),
    ('👥 Usuarios', 'Login/Logout con JWT', '✅'),
    ('👥 Usuarios', 'Roles (admin/usuario)', '✅'),
    ('📚 Libros', 'CRUD completo de libros', '✅'),
    ('📚 Libros', 'Búsqueda por título/autor', '✅'),
    ('📚 Libros', 'Filtrado por categoría', '✅'),
    ('📚 Libros', 'Ordenamiento múltiple', '✅'),
    ('📖 Lectura', 'Lector de PDF integrado', '✅'),
    ('📖 Lectura', 'Lector de texto con zoom', '✅'),
    ('📖 Lectura', 'Contador de vistas', '✅'),
    ('🖥️ PC Remoto', 'Conexión vía Cloudflare', '✅'),
    ('🖥️ PC Remoto', 'Servir PDFs locales', '✅'),
    ('🖥️ PC Remoto', 'Servir portadas locales', '✅'),
    ('🤖 Chatbot', 'Asistente Botpress', '✅'),
    ('💳 Pagos', 'Donaciones con PayPal', '✅'),
]

for i, (c1, c2, c3) in enumerate(functionalities):
    row = func_table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, '8b5cf6')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True
    elif i % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'F5F3FF')
    # Color verde para el check
    if c3 == '✅':
        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(16, 185, 129)

doc.add_page_break()

# ============= 4. LIMITACIONES =============
h = doc.add_heading('4. ❌ Limitaciones (Lo que NO puede hacer)', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(239, 68, 68)

# Tabla de limitaciones
lim_table = doc.add_table(rows=11, cols=3)
lim_table.style = 'Table Grid'

limitations = [
    ('Área', 'Limitación', 'Alternativa'),
    ('👥 Usuarios', 'Recuperar contraseña', 'Contactar admin'),
    ('👥 Usuarios', 'Editar perfil propio', 'Solo admin puede'),
    ('👥 Usuarios', 'Foto de perfil', 'No disponible'),
    ('📚 Libros', 'Subir archivos directamente', 'Usar URLs externas'),
    ('📚 Libros', 'Formato EPUB', 'Solo PDF y texto'),
    ('📚 Libros', 'Anotaciones/marcadores', 'No disponible'),
    ('🖥️ PC Remoto', 'Múltiples PCs', 'Solo 1 PC a la vez'),
    ('🖥️ PC Remoto', 'Sincronización automática', 'Manual'),
    ('💬 Chat', 'Chat entre usuarios', 'Solo chatbot'),
    ('📱 Móvil', 'App nativa', 'Solo web responsive'),
]

for i, (c1, c2, c3) in enumerate(limitations):
    row = lim_table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'ef4444')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True
    elif i % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'FEF2F2')

doc.add_paragraph()

# Invitados
doc.add_heading('Restricciones para Invitados (sin cuenta)', level=2)
inv_table = doc.add_table(rows=1, cols=1)
cell = inv_table.rows[0].cells[0]
set_cell_shading(cell, 'FEF3C7')
cell.text = """⚠️ Los usuarios sin cuenta solo pueden leer 1 libro.
Después deben registrarse para continuar."""

doc.add_page_break()

# ============= 5. MAPA DE DEPENDENCIAS =============
h = doc.add_heading('5. 🗺️ Mapa de Dependencias de Archivos', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(245, 158, 11)

doc.add_paragraph('Este mapa muestra qué archivos necesita cada página para funcionar correctamente.')

# PÁGINA PRINCIPAL
doc.add_heading('📄 index.html (Página Principal)', level=2)
idx_table = doc.add_table(rows=6, cols=2)
idx_table.style = 'Table Grid'

idx_deps = [
    ('Tipo', 'Archivo'),
    ('🎨 CSS', '/css/index.css'),
    ('📜 JS', '/javascript/index.js'),
    ('🔌 API', '/api/books/featured, /api/books/recent'),
    ('🔌 API', '/api/categories, /api/auth/me'),
    ('🤖 Externo', 'Botpress (cdn.botpress.cloud)'),
]

for i, (t, f) in enumerate(idx_deps):
    row = idx_table.rows[i]
    row.cells[0].text = t
    row.cells[1].text = f
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'f59e0b')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()

# CATÁLOGO
doc.add_heading('📄 catalog.html (Catálogo)', level=2)
cat_table = doc.add_table(rows=5, cols=2)
cat_table.style = 'Table Grid'

cat_deps = [
    ('Tipo', 'Archivo'),
    ('🎨 CSS', '/css/index.css, /css/catalog.css'),
    ('📜 JS', 'Inline (dentro del HTML)'),
    ('🔌 API', '/api/books, /api/categories'),
    ('🤖 Externo', 'Botpress'),
]

for i, (t, f) in enumerate(cat_deps):
    row = cat_table.rows[i]
    row.cells[0].text = t
    row.cells[1].text = f
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'f59e0b')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()

# DASHBOARD
doc.add_heading('📄 dashboard.html (Panel Admin)', level=2)
dash_table = doc.add_table(rows=6, cols=2)
dash_table.style = 'Table Grid'

dash_deps = [
    ('Tipo', 'Archivo'),
    ('🎨 CSS', '/css/index.css, /css/dashboard.css'),
    ('📜 JS', '/javascript/dashboard.js'),
    ('🔌 API', '/api/stats, /api/books, /api/categories'),
    ('🔌 API', '/api/users, /api/remote/status, /api/remote/files'),
    ('🤖 Externo', 'Botpress'),
]

for i, (t, f) in enumerate(dash_deps):
    row = dash_table.rows[i]
    row.cells[0].text = t
    row.cells[1].text = f
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'f59e0b')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()

# READER
doc.add_heading('📄 reader.html (Lector)', level=2)
read_table = doc.add_table(rows=4, cols=2)
read_table.style = 'Table Grid'

read_deps = [
    ('Tipo', 'Archivo'),
    ('🎨 CSS', 'Inline (estilos en el HTML)'),
    ('📜 JS', 'Inline (script en el HTML)'),
    ('🔌 API', '/api/books/:id, /api/books/:id/read'),
]

for i, (t, f) in enumerate(read_deps):
    row = read_table.rows[i]
    row.cells[0].text = t
    row.cells[1].text = f
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'f59e0b')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()

# LOGIN/REGISTER
doc.add_heading('📄 login.html / register.html', level=2)
log_table = doc.add_table(rows=4, cols=2)
log_table.style = 'Table Grid'

log_deps = [
    ('Tipo', 'Archivo'),
    ('🎨 CSS', '/css/login.css'),
    ('📜 JS', '/javascript/login.js o /javascript/register.js'),
    ('🔌 API', '/api/auth/login o /api/auth/register'),
]

for i, (t, f) in enumerate(log_deps):
    row = log_table.rows[i]
    row.cells[0].text = t
    row.cells[1].text = f
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'f59e0b')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_page_break()

# Diagrama visual de dependencias
doc.add_heading('Diagrama de Dependencias', level=2)

dep_diagram = doc.add_table(rows=1, cols=1)
dep_diagram.style = 'Table Grid'
cell = dep_diagram.rows[0].cells[0]
set_cell_shading(cell, 'FFFBEB')

diagram2 = '''
                           ┌─────────────────┐
                           │   server.js     │ ◄── Punto de entrada
                           └────────┬────────┘
                                    │
            ┌───────────────────────┼───────────────────────┐
            │                       │                       │
            ▼                       ▼                       ▼
    ┌───────────────┐      ┌───────────────┐      ┌───────────────┐
    │    MODELOS    │      │     RUTAS     │      │    PÚBLICO    │
    │   /models/    │      │   /routes/    │      │   /public/    │
    └───────┬───────┘      └───────┬───────┘      └───────┬───────┘
            │                      │                      │
    ┌───────┴───────┐      ┌───────┴───────┐      ┌───────┴───────┐
    │ • User.js     │      │ • auth.js     │      │ /html/        │
    │ • Book.js     │      │ • books.js    │      │ • index.html  │
    │ • Category.js │      │ • categories  │      │ • catalog.html│
    │ • Remote...   │      │ • users.js    │      │ • dashboard   │
    │ • AccessLog   │      │ • remote.js   │      │ • reader.html │
    │ • Donation    │      │ • paypal.js   │      │ • login.html  │
    └───────────────┘      └───────────────┘      └───────┬───────┘
                                                         │
                                                 ┌───────┴───────┐
                                                 │ /css/         │
                                                 │ • index.css   │
                                                 │ • catalog.css │
                                                 │ • dashboard   │
                                                 │ • book.css    │
                                                 │ • login.css   │
                                                 └───────┬───────┘
                                                         │
                                                 ┌───────┴───────┐
                                                 │ /javascript/  │
                                                 │ • index.js    │
                                                 │ • dashboard.js│
                                                 │ • login.js    │
                                                 │ • register.js │
                                                 └───────────────┘
'''
cell.text = diagram2
for para in cell.paragraphs:
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

doc.add_page_break()

# ============= 6. FLUJOS DE DATOS =============
h = doc.add_heading('6. 🔄 Flujos de Datos', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(6, 182, 212)

# Flujo de autenticación
doc.add_heading('Flujo de Autenticación', level=2)

auth_flow = doc.add_table(rows=1, cols=1)
auth_flow.style = 'Table Grid'
cell = auth_flow.rows[0].cells[0]
set_cell_shading(cell, 'ECFEFF')

auth_diagram = '''
    ┌──────────┐         ┌──────────┐         ┌──────────┐
    │ Usuario  │         │ Backend  │         │ MongoDB  │
    └────┬─────┘         └────┬─────┘         └────┬─────┘
         │                    │                    │
         │  POST /login       │                    │
         │  {email, password} │                    │
         │───────────────────>│                    │
         │                    │  findOne({email})  │
         │                    │───────────────────>│
         │                    │                    │
         │                    │     user data      │
         │                    │<───────────────────│
         │                    │                    │
         │                    │ bcrypt.compare()   │
         │                    │                    │
         │                    │ jwt.sign()         │
         │                    │                    │
         │  Set-Cookie: token │                    │
         │  {user}            │                    │
         │<───────────────────│                    │
         │                    │                    │
'''
cell.text = auth_diagram
for para in cell.paragraphs:
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()

# Flujo de PC Remoto
doc.add_heading('Flujo de PC Remoto', level=2)

remote_flow = doc.add_table(rows=1, cols=1)
remote_flow.style = 'Table Grid'
cell = remote_flow.rows[0].cells[0]
set_cell_shading(cell, 'F0FDFA')

remote_diagram = '''
    ┌──────────┐     ┌───────────┐     ┌──────────┐     ┌──────────┐
    │ Navegador│     │ Biblioteca│     │ MongoDB  │     │ PC Remoto│
    └────┬─────┘     └─────┬─────┘     └────┬─────┘     └────┬─────┘
         │                 │                │                │
         │                 │                │   /register    │
         │                 │                │<───────────────│
         │                 │                │                │
         │                 │  save URL      │                │
         │                 │───────────────>│                │
         │                 │                │                │
         │ GET /book/123   │                │                │
         │ (cover: remote:x.jpg)            │                │
         │────────────────>│                │                │
         │                 │                │                │
         │                 │  get PC URL    │                │
         │                 │───────────────>│                │
         │                 │                │                │
         │                 │                │  GET /file/... │
         │                 │                │───────────────>│
         │                 │                │                │
         │                 │                │   image data   │
         │                 │<──────────────────────────────│
         │   image         │                │                │
         │<────────────────│                │                │
         │                 │                │                │
'''
cell.text = remote_diagram
for para in cell.paragraphs:
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_page_break()

# ============= 7. BASE DE DATOS =============
h = doc.add_heading('7. 🗄️ Diagrama de Base de Datos', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(236, 72, 153)

db_diagram = doc.add_table(rows=1, cols=1)
db_diagram.style = 'Table Grid'
cell = db_diagram.rows[0].cells[0]
set_cell_shading(cell, 'FDF2F8')

db = '''
┌─────────────────────────────────────────────────────────────────────────┐
│                           MONGODB - biblioteca_virtual                   │
└─────────────────────────────────────────────────────────────────────────┘

┌─────────────────────┐     ┌─────────────────────┐     ┌─────────────────┐
│       users         │     │       books         │     │   categories    │
├─────────────────────┤     ├─────────────────────┤     ├─────────────────┤
│ _id: ObjectId       │     │ _id: ObjectId       │     │ _id: ObjectId   │
│ name: String ✱      │     │ title: String ✱     │     │ name: String ✱  │
│ email: String ✱ 🔑  │     │ author: String ✱    │     │ slug: String ✱  │
│ passwordHash: String│     │ description: String │     │ description: Str│
│ role: \"user\"|\"admin\"│     │ categories: [String]│─ ─ >│                 │
│ createdAt: Date     │     │ coverUrl: String    │     └─────────────────┘
└─────────────────────┘     │ pdfUrl: String      │
                            │ content: String     │
                            │ views: Number       │
                            │ createdAt: Date     │
                            └─────────────────────┘

┌─────────────────────┐     ┌─────────────────────┐     ┌─────────────────┐
│   remotesources     │     │     accesslogs      │     │    donations    │
├─────────────────────┤     ├─────────────────────┤     ├─────────────────┤
│ _id: ObjectId       │     │ _id: ObjectId       │     │ _id: ObjectId   │
│ name: String        │     │ identifier: String  │     │ amount: Number  │
│ url: String ✱       │     │ bookId: String      │     │ currency: String│
│ apiKey: String ✱    │     │ createdAt: Date     │     │ paypalId: String│
│ isOnline: Boolean   │     └─────────────────────┘     │ status: String  │
│ lastSeen: Date      │                                 │ createdAt: Date │
│ createdAt: Date     │     ✱ = Requerido               └─────────────────┘
└─────────────────────┘     🔑 = Único
'''
cell.text = db
for para in cell.paragraphs:
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_page_break()

# ============= 8. GUÍA DE DESPLIEGUE =============
h = doc.add_heading('8. 🚀 Guía Rápida de Despliegue', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(99, 102, 241)

# Pasos
doc.add_heading('Pasos para Iniciar', level=2)

steps_table = doc.add_table(rows=6, cols=2)
steps_table.style = 'Table Grid'

steps = [
    ('Paso', 'Comando / Acción'),
    ('1️⃣ Instalar dependencias', 'npm install'),
    ('2️⃣ Configurar .env', 'Copiar .env.example a .env y editar'),
    ('3️⃣ Iniciar servidor', 'npm start'),
    ('4️⃣ Crear datos iniciales', 'POST /api/seed (una sola vez)'),
    ('5️⃣ Acceder', 'http://localhost:3000'),
]

for i, (step, cmd) in enumerate(steps):
    row = steps_table.rows[i]
    row.cells[0].text = step
    row.cells[1].text = cmd
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, '6366f1')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()

# Variables de entorno
doc.add_heading('Variables de Entorno Requeridas', level=2)

env_table = doc.add_table(rows=6, cols=3)
env_table.style = 'Table Grid'

envs = [
    ('Variable', 'Descripción', 'Ejemplo'),
    ('MONGO_URL', 'URL de MongoDB', 'mongodb+srv://...'),
    ('DB_NAME', 'Nombre de la BD', 'biblioteca_virtual'),
    ('JWT_SECRET', 'Clave para tokens', 'mi_clave_secreta'),
    ('PORT', 'Puerto del servidor', '3000'),
    ('REMOTE_API_KEY', 'Clave para PC remoto', 'MiBiblioteca2024...'),
]

for i, (v, d, e) in enumerate(envs):
    row = env_table.rows[i]
    row.cells[0].text = v
    row.cells[1].text = d
    row.cells[2].text = e
    if i == 0:
        for cell in row.cells:
            set_cell_shading(cell, '4F46E5')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# Credenciales por defecto
doc.add_heading('Credenciales por Defecto', level=2)

cred_table = doc.add_table(rows=1, cols=1)
cell = cred_table.rows[0].cells[0]
set_cell_shading(cell, 'EEF2FF')
cell.text = '👤 Admin: admin@biblioteca.com / admin123'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Guardar documento
doc.save('docs/Proyecto_Completo.docx')
print('✅ Documento del Proyecto Completo creado')
