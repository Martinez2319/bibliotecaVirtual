from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================
# MANUAL TÉCNICO
# =============================================
doc = Document()

# Título
title = doc.add_heading('Biblioteca Virtual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Manual Técnico')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Índice
doc.add_heading('Índice', level=1)
doc.add_paragraph('1. Arquitectura del Sistema')
doc.add_paragraph('2. Tecnologías Utilizadas')
doc.add_paragraph('3. Estructura de Archivos')
doc.add_paragraph('4. Base de Datos')
doc.add_paragraph('5. API REST - Endpoints')
doc.add_paragraph('6. Autenticación')
doc.add_paragraph('7. Sistema de PC Remoto')
doc.add_paragraph('8. Instalación y Despliegue')
doc.add_paragraph('9. Variables de Entorno')
doc.add_paragraph('10. Solución de Problemas')

doc.add_page_break()

# 1. Arquitectura
doc.add_heading('1. Arquitectura del Sistema', level=1)
doc.add_paragraph('El sistema sigue una arquitectura cliente-servidor con los siguientes componentes:')
doc.add_paragraph('')
doc.add_paragraph('┌─────────────────┐     ┌─────────────────┐     ┌─────────────────┐')
doc.add_paragraph('│   FRONTEND      │────▶│    BACKEND      │────▶│   MONGODB       │')
doc.add_paragraph('│   (HTML/CSS/JS) │     │   (Express.js)  │     │   (Atlas/Local) │')
doc.add_paragraph('└─────────────────┘     └─────────────────┘     └─────────────────┘')
doc.add_paragraph('                               │')
doc.add_paragraph('                               ▼')
doc.add_paragraph('                        ┌─────────────────┐')
doc.add_paragraph('                        │   PC REMOTO     │')
doc.add_paragraph('                        │   (Express.js)  │')
doc.add_paragraph('                        └─────────────────┘')

doc.add_heading('Flujo de datos:', level=2)
doc.add_paragraph('1. El usuario accede al frontend (páginas HTML estáticas)')
doc.add_paragraph('2. JavaScript hace peticiones a la API REST (/api/*)')
doc.add_paragraph('3. El backend procesa las peticiones y consulta MongoDB')
doc.add_paragraph('4. Para archivos remotos, el backend hace proxy al PC remoto')

# 2. Tecnologías
doc.add_heading('2. Tecnologías Utilizadas', level=1)

doc.add_heading('Backend:', level=2)
doc.add_paragraph('• Node.js v18+')
doc.add_paragraph('• Express.js 4.x - Framework web')
doc.add_paragraph('• Mongoose 8.x - ODM para MongoDB')
doc.add_paragraph('• JWT (jsonwebtoken) - Autenticación')
doc.add_paragraph('• bcryptjs - Hash de contraseñas')
doc.add_paragraph('• cookie-parser - Manejo de cookies')
doc.add_paragraph('• cors - Cross-Origin Resource Sharing')
doc.add_paragraph('• axios - Cliente HTTP (para PC remoto)')

doc.add_heading('Frontend:', level=2)
doc.add_paragraph('• HTML5')
doc.add_paragraph('• CSS3 (responsive, flexbox, grid)')
doc.add_paragraph('• JavaScript vanilla (ES6+)')
doc.add_paragraph('• Botpress - Chatbot')

doc.add_heading('Base de datos:', level=2)
doc.add_paragraph('• MongoDB Atlas (cloud) o MongoDB local')

doc.add_heading('Infraestructura:', level=2)
doc.add_paragraph('• Cloudflare Tunnels - Para exponer servidores locales')

# 3. Estructura de Archivos
doc.add_heading('3. Estructura de Archivos', level=1)
doc.add_paragraph('biblioteca-virtual/')
doc.add_paragraph('├── server.js              # Servidor principal Express')
doc.add_paragraph('├── package.json           # Dependencias del proyecto')
doc.add_paragraph('├── .env                   # Variables de entorno')
doc.add_paragraph('├── models/                # Modelos de MongoDB')
doc.add_paragraph('│   ├── User.js           # Modelo de usuario')
doc.add_paragraph('│   ├── Book.js           # Modelo de libro')
doc.add_paragraph('│   ├── Category.js       # Modelo de categoría')
doc.add_paragraph('│   ├── RemoteSource.js   # Modelo de PC remoto')
doc.add_paragraph('│   ├── AccessLog.js      # Control de acceso invitados')
doc.add_paragraph('│   └── Donation.js       # Donaciones PayPal')
doc.add_paragraph('├── routes/                # Rutas de la API')
doc.add_paragraph('│   ├── auth.js           # Autenticación')
doc.add_paragraph('│   ├── books.js          # CRUD libros')
doc.add_paragraph('│   ├── categories.js     # CRUD categorías')
doc.add_paragraph('│   ├── users.js          # Gestión usuarios')
doc.add_paragraph('│   ├── remote.js         # PC remoto')
doc.add_paragraph('│   └── paypal.js         # Donaciones')
doc.add_paragraph('├── public/                # Archivos estáticos')
doc.add_paragraph('│   ├── html/             # Páginas HTML')
doc.add_paragraph('│   ├── css/              # Estilos')
doc.add_paragraph('│   └── javascript/       # Scripts del frontend')
doc.add_paragraph('└── pc-server/             # Servidor para PC remoto')
doc.add_paragraph('    ├── server.js         # Servidor Express')
doc.add_paragraph('    ├── config.json       # Configuración')
doc.add_paragraph('    └── iniciar.bat       # Script de inicio')

# 4. Base de Datos
doc.add_heading('4. Base de Datos', level=1)

doc.add_heading('Colección: users', level=2)
doc.add_paragraph('• name: String (requerido)')
doc.add_paragraph('• email: String (único, requerido)')
doc.add_paragraph('• passwordHash: String (requerido)')
doc.add_paragraph('• role: String (user | admin, default: user)')
doc.add_paragraph('• createdAt: Date')

doc.add_heading('Colección: books', level=2)
doc.add_paragraph('• title: String (requerido)')
doc.add_paragraph('• author: String (requerido)')
doc.add_paragraph('• description: String')
doc.add_paragraph('• categories: [String]')
doc.add_paragraph('• coverUrl: String')
doc.add_paragraph('• pdfUrl: String')
doc.add_paragraph('• content: String')
doc.add_paragraph('• views: Number (default: 0)')
doc.add_paragraph('• createdAt: Date')

doc.add_heading('Colección: categories', level=2)
doc.add_paragraph('• name: String (requerido)')
doc.add_paragraph('• slug: String (único, requerido)')
doc.add_paragraph('• description: String')

doc.add_heading('Colección: remotesources', level=2)
doc.add_paragraph('• name: String')
doc.add_paragraph('• url: String (requerido)')
doc.add_paragraph('• apiKey: String (requerido)')
doc.add_paragraph('• isOnline: Boolean')
doc.add_paragraph('• lastSeen: Date')

doc.add_heading('Colección: accesslogs', level=2)
doc.add_paragraph('• identifier: String (IP del usuario)')
doc.add_paragraph('• bookId: String')
doc.add_paragraph('• createdAt: Date')

# 5. API REST
doc.add_heading('5. API REST - Endpoints', level=1)

doc.add_heading('Autenticación (/api/auth)', level=2)
doc.add_paragraph('POST /api/auth/register - Registrar usuario')
doc.add_paragraph('POST /api/auth/login - Iniciar sesión')
doc.add_paragraph('POST /api/auth/logout - Cerrar sesión')
doc.add_paragraph('GET /api/auth/me - Obtener usuario actual')

doc.add_heading('Libros (/api/books)', level=2)
doc.add_paragraph('GET /api/books - Listar libros (query: search, category, sort)')
doc.add_paragraph('GET /api/books/featured - Libros destacados')
doc.add_paragraph('GET /api/books/recent - Libros recientes')
doc.add_paragraph('GET /api/books/:id - Obtener libro')
doc.add_paragraph('POST /api/books - Crear libro (admin)')
doc.add_paragraph('PUT /api/books/:id - Actualizar libro (admin)')
doc.add_paragraph('DELETE /api/books/:id - Eliminar libro (admin)')
doc.add_paragraph('POST /api/books/:id/read - Registrar lectura')

doc.add_heading('Categorías (/api/categories)', level=2)
doc.add_paragraph('GET /api/categories - Listar categorías')
doc.add_paragraph('POST /api/categories - Crear categoría (admin)')
doc.add_paragraph('PUT /api/categories/:id - Actualizar categoría (admin)')
doc.add_paragraph('DELETE /api/categories/:id - Eliminar categoría (admin)')

doc.add_heading('Usuarios (/api/users)', level=2)
doc.add_paragraph('GET /api/users - Listar usuarios (admin)')
doc.add_paragraph('PUT /api/users/:id - Actualizar rol (admin)')
doc.add_paragraph('DELETE /api/users/:id - Eliminar usuario (admin)')

doc.add_heading('PC Remoto (/api/remote)', level=2)
doc.add_paragraph('POST /api/remote/register - Registrar PC (requiere API key)')
doc.add_paragraph('POST /api/remote/heartbeat - Mantener conexión')
doc.add_paragraph('POST /api/remote/disconnect - Desconectar PC')
doc.add_paragraph('GET /api/remote/status - Estado del PC (admin)')
doc.add_paragraph('GET /api/remote/files - Listar archivos (admin)')
doc.add_paragraph('GET /api/remote/file/:type/:filename - Obtener archivo')

# 6. Autenticación
doc.add_heading('6. Autenticación', level=1)
doc.add_paragraph('El sistema usa JWT (JSON Web Tokens) almacenados en cookies HTTP-only.')

doc.add_heading('Flujo de autenticación:', level=2)
doc.add_paragraph('1. Usuario envía credenciales a POST /api/auth/login')
doc.add_paragraph('2. Backend verifica contraseña con bcrypt')
doc.add_paragraph('3. Si es válida, genera JWT con el ID del usuario')
doc.add_paragraph('4. JWT se almacena en cookie \"token\"')
doc.add_paragraph('5. En cada petición, middleware extrae y verifica el JWT')
doc.add_paragraph('6. Si es válido, req.user contiene los datos del usuario')

doc.add_heading('Middleware de autenticación:', level=2)
doc.add_paragraph('Ubicado en server.js, se ejecuta en cada petición:')
doc.add_paragraph('• Extrae token de cookies o header Authorization')
doc.add_paragraph('• Verifica firma con JWT_SECRET')
doc.add_paragraph('• Busca usuario en MongoDB')
doc.add_paragraph('• Asigna a req.user')

# 7. Sistema PC Remoto
doc.add_heading('7. Sistema de PC Remoto', level=1)

doc.add_heading('Arquitectura:', level=2)
doc.add_paragraph('┌──────────────┐    Cloudflare    ┌──────────────┐')
doc.add_paragraph('│  PC Usuario  │◀────Tunnel────▶│  Biblioteca  │')
doc.add_paragraph('│  (pc-server) │                 │  (server.js) │')
doc.add_paragraph('│  :3001       │                 │  :3000       │')
doc.add_paragraph('└──────────────┘                 └──────────────┘')

doc.add_heading('Comunicación:', level=2)
doc.add_paragraph('1. PC inicia y obtiene URL de Cloudflare')
doc.add_paragraph('2. PC se registra con POST /api/remote/register')
doc.add_paragraph('3. Envía heartbeat cada 30 segundos')
doc.add_paragraph('4. Biblioteca guarda URL y estado en MongoDB')
doc.add_paragraph('5. Al solicitar archivo remote:*, biblioteca hace proxy al PC')

doc.add_heading('Seguridad:', level=2)
doc.add_paragraph('• API Key compartida entre PC y servidor')
doc.add_paragraph('• Validación de paths para evitar directory traversal')
doc.add_paragraph('• Timeout de conexión de 60 segundos')

# 8. Instalación
doc.add_heading('8. Instalación y Despliegue', level=1)

doc.add_heading('Requisitos:', level=2)
doc.add_paragraph('• Node.js 18 o superior')
doc.add_paragraph('• MongoDB (local o Atlas)')
doc.add_paragraph('• npm o yarn')

doc.add_heading('Pasos:', level=2)
doc.add_paragraph('1. Clonar repositorio')
doc.add_paragraph('2. npm install')
doc.add_paragraph('3. Crear archivo .env con las variables')
doc.add_paragraph('4. npm start')

doc.add_heading('Para PC remoto:', level=2)
doc.add_paragraph('1. cd pc-server')
doc.add_paragraph('2. npm install')
doc.add_paragraph('3. Instalar cloudflared')
doc.add_paragraph('4. Ejecutar iniciar.bat')

# 9. Variables de Entorno
doc.add_heading('9. Variables de Entorno', level=1)
doc.add_paragraph('Archivo .env en la raíz del proyecto:')
doc.add_paragraph('')
doc.add_paragraph('MONGO_URL=mongodb+srv://...')
doc.add_paragraph('DB_NAME=biblioteca_virtual')
doc.add_paragraph('JWT_SECRET=clave_secreta_jwt')
doc.add_paragraph('PORT=3000')
doc.add_paragraph('REMOTE_API_KEY=clave_para_pc_remoto')
doc.add_paragraph('PAYPAL_CLIENT_ID=id_de_paypal')
doc.add_paragraph('PAYPAL_SECRET=secreto_de_paypal')
doc.add_paragraph('PAYPAL_MODE=sandbox')

# 10. Solución de Problemas
doc.add_heading('10. Solución de Problemas', level=1)

doc.add_heading('Error de conexión MongoDB:', level=2)
doc.add_paragraph('• Verificar MONGO_URL en .env')
doc.add_paragraph('• Verificar que no tenga el nombre de DB duplicado')
doc.add_paragraph('• Verificar IP en whitelist de Atlas')

doc.add_heading('PC remoto no conecta:', level=2)
doc.add_paragraph('• Verificar que ambos túneles Cloudflare estén activos')
doc.add_paragraph('• Verificar que API_KEY coincida')
doc.add_paragraph('• Revisar logs del servidor')

doc.add_heading('JWT inválido:', level=2)
doc.add_paragraph('• Verificar JWT_SECRET en .env')
doc.add_paragraph('• Limpiar cookies del navegador')

doc.save('docs/Manual_Tecnico.docx')
print('✅ Manual Técnico creado')
